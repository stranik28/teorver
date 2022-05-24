from crypt import methods
from fileinput import filename

from docxcompose.composer import Composer
from docx import Document
import random
import shutil
import os
from flask import Flask, render_template, send_from_directory, request


app = Flask(__name__)


@app.route("/", methods=["GET"])
def variants():
    return render_template("index.html")

@app.route("/get_var", methods=["POST"])
def ret():
    n = int(request.form.get("numb"))
    shutil.rmtree("done")
    os.mkdir("done")
    for i in range(1,n+1):
        for j in range(1,11):
            r = str(random.randint(0, 29))
            master = Document('1/'+r+'.docx')
            composer = Composer(master)
            r = str(random.randint(0, 29))
            doc1 = Document('2/'+r+'.docx')
            r = str(random.randint(0, 29))
            doc2 = Document('3/'+r+'.docx')
            r = str(random.randint(0, 29))
            doc3 = Document('4/'+r+'.docx')
            r = str(random.randint(0, 29))
            doc4 = Document('5/'+r+'.docx')
            r = str(random.randint(0, 29))
            doc5 = Document('6/'+r+'.docx')
            r = str(random.randint(0, 29))
            doc6 = Document('7/'+r+'.docx')
            r = str(random.randint(0, 29))
            doc7 = Document('8/'+r+'.docx')
            r = str(random.randint(0, 29))
            doc8 = Document('9/'+r+'.docx')
            r = str(random.randint(0, 29))
            doc9 = Document('10/'+r+'.docx')
            composer.append(doc1)
            composer.append(doc2)
            composer.append(doc3)
            composer.append(doc4)
            composer.append(doc5)
            composer.append(doc6)
            composer.append(doc7)
            composer.append(doc8)
            composer.append(doc9)
            composer.save("done/" + str(i) + ".docx")
    shutil.make_archive("variants", 'zip', "done")
    return send_from_directory(directory = "", filename = "variants.zip", as_attachment = True)

    


app.run()